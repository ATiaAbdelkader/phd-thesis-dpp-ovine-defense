#!/usr/bin/env python3
"""Extract text content from all DOCX files in upload folder."""
import os
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx -q")
    from docx import Document

UPLOAD_DIR = Path("/home/z/my-project/upload")
OUT_DIR = Path("/home/z/my-project/extracted")
OUT_DIR.mkdir(parents=True, exist_ok=True)

docx_files = sorted(UPLOAD_DIR.glob("*.docx"))
print(f"Found {len(docx_files)} DOCX files")

for docx_path in docx_files:
    print(f"\n{'='*70}")
    print(f"Processing: {docx_path.name}")
    print(f"{'='*70}")
    try:
        doc = Document(str(docx_path))
        out_lines = []
        # Iterate paragraphs and tables in document order
        # python-docx doesn't easily give order; capture paragraphs and tables separately
        n_paragraphs = len(doc.paragraphs)
        n_tables = len(doc.tables)
        print(f"  Paragraphs: {n_paragraphs}, Tables: {n_tables}")

        out_lines.append(f"# FILE: {docx_path.name}\n")
        out_lines.append(f"Paragraphs: {n_paragraphs} | Tables: {n_tables}\n")
        out_lines.append("=" * 70 + "\n")

        # Paragraphs
        for i, p in enumerate(doc.paragraphs):
            style = p.style.name if p.style else "Normal"
            text = p.text.strip()
            if text:
                out_lines.append(f"[P{i:04d} | {style}] {text}")
            else:
                out_lines.append(f"[P{i:04d} | {style}] (empty)")

        # Tables
        out_lines.append("\n" + "=" * 70)
        out_lines.append("TABLES")
        out_lines.append("=" * 70)
        for ti, table in enumerate(doc.tables):
            out_lines.append(f"\n--- Table {ti+1} ({len(table.rows)} rows × {len(table.columns)} cols) ---")
            for ri, row in enumerate(table.rows):
                cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
                out_lines.append(f"  R{ri:02d}: " + " || ".join(cells))

        out_path = OUT_DIR / (docx_path.stem + ".txt")
        out_path.write_text("\n".join(out_lines), encoding="utf-8")
        print(f"  -> Saved to {out_path}")
    except Exception as e:
        print(f"  ERROR: {e}")

print("\n\nDone. All extracted to /home/z/my-project/extracted/")
